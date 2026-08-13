import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

document = Document()

# Set Margins for all sections
for section in document.sections:
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1.5)

# Set Default Style (Normal)
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.line_spacing = 1.5
paragraph_format.space_after = Pt(9)

def add_chapter(doc, num, title, content_paragraphs):
    if num not in ["Abstract", "References", "Annexures"]:
        # Add Chapter - X aligned right
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"Chapter - {num}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    
    # Title centered, Font 14, All Caps
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(title.upper())
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(14)
    run_title.bold = True
    
    # Add paragraphs
    for text in content_paragraphs:
        if text.startswith("### "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text[4:])
            r.bold = True
        elif text.startswith("[PLACEHOLDER]"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.italic = True
        else:
            doc.add_paragraph(text)
            
    doc.add_page_break()

# Title page details (as requested in Admin Details)
p = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("SUMMER INTERNSHIP REPORT\n").bold = True
p.add_run("on\n")
run_title = p.add_run("DEVELOPMENT OF ENTERPRISE LMS PORTAL AT XEBIA\n")
run_title.font.size = Pt(14)
run_title.bold = True
p.add_run("\nSubmitted by:\nKESHAVBHATT KUMARI ANAGHA\n")
p.add_run("University: Sir Padampat Singhania University, Udaipur\n")
p.add_run("Department of Computing and Informatics, School of Engineering & Sciences\n\n")
p.add_run("Company: Xebia IT Architects Pvt Ltd. (Domain: Product Engineering)\n")
p.add_run("Duration: 2 Months (2 June 2026 – 31 July 2026)\n\n")
p.add_run("External Guide: Biplove Singh\n")
p.add_run("Internal Guide: Dr. Chandani Joshi\n")

document.add_page_break()

# Content Sections
abstract = [
    "This report details the backend architecture and development of an enterprise-grade Learning Management System (LMS) at Xebia IT Architects Pvt Ltd. over a two-month summer internship. The project aimed to build a scalable, modern, and headless LMS to serve enterprise clients. My role was focused entirely on the backend engineering and database architecture, utilizing a robust technology stack comprising Java 17, Spring Boot 3.2.5, PostgreSQL, Flyway, and MapStruct. Working within a specialized five-member backend team, we adopted an iterative development approach. We successfully engineered and deployed a series of RESTful APIs to manage the core curriculum (Categories, Courses, Modules, Submodules, Content), followed by comprehensive modules for Assessments, Events, Event Registrations, and Batches & Student Management. A key aspect of the architecture was the implementation of a multi-tenant data isolation strategy to securely segregate client data. Rigorous API testing via Postman, combined with manual database verification and concurrency load testing, ensured system reliability and scalability under load. The internship culminated in the successful integration of our backend APIs with the frontend team's React application, followed by live deployment. This experience provided profound insights into agile software development, enterprise database management, and building high-performance server-side applications."
]
add_chapter(document, "Abstract", "Abstract", abstract)

ch1 = [
    "Xebia IT Architects Pvt Ltd. is a leading global IT consultancy and software development firm, recognized for its expertise in agile methodologies, DevOps, and Product Engineering. The Product Engineering domain within Xebia focuses on building robust, scalable, and innovative digital products tailored to complex enterprise requirements.",
    "The primary objective of this project was to design and develop a scalable, highly available enterprise Learning Management System (LMS). As digital learning environments become increasingly sophisticated, organizations require flexible and headless LMS solutions that can seamlessly integrate with various client-facing interfaces.",
    "My specific scope of work was strictly confined to backend and database architecture. I was responsible for designing relational database schemas, creating efficient RESTful APIs, and ensuring secure data transactions. I was entirely excluded from the frontend User Interface (UI) development, allowing for a concentrated effort on backend logic, data validation, and service optimization."
]
add_chapter(document, "1", "Introduction", ch1)

ch2 = [
    "The architecture of Learning Management Systems has evolved significantly over the past decade. Traditional LMS platforms were typically built as monolithic applications, tightly coupling the backend logic with the frontend presentation layer. While simpler to deploy initially, monoliths suffer from scalability bottlenecks and limit the flexibility to adopt new frontend frameworks.",
    "Modern enterprise applications have transitioned towards headless, API-driven architectures [1]. In a headless LMS, the backend content repository and logic are completely separated from the frontend presentation layer, communicating exclusively via RESTful or GraphQL APIs. This separation of concerns allows for parallel development, greater technological flexibility, and enhanced scalability.",
    "The use of microservices and component-based architectures using frameworks like Spring Boot has become the industry standard for enterprise backend systems [2]. Spring Boot simplifies the bootstrapping and development of new Spring applications, providing robust dependency injection, embedded web servers, and comprehensive security features out-of-the-box [3]. Furthermore, managing database schema evolutions is a critical challenge in agile development. Tools like Flyway have emerged as essential mechanisms for automated and version-controlled database migrations [4]."
]
add_chapter(document, "2", "Literature Review", ch2)

ch3 = [
    "Xebia IT Architects Pvt Ltd. operates at the forefront of digital transformation, emphasizing high-quality software craftsmanship. Within the Product Engineering division, teams follow agile practices, emphasizing continuous integration, test-driven development, and iterative delivery.",
    "The development of the LMS backend was powered by a modern, enterprise-grade technology stack:",
    "### Java 17",
    "Java 17 was chosen as the core programming language due to its Long-Term Support (LTS) status, performance enhancements, and modern features like records and enhanced switch statements, which significantly reduce boilerplate code.",
    "### Spring Boot 3.2.5",
    "The Spring Boot framework provided the foundation for building the RESTful APIs. Its auto-configuration and extensive ecosystem enabled rapid development of web services, security configurations, and data access layers.",
    "### PostgreSQL",
    "PostgreSQL, a highly advanced open-source relational database, was selected for its robustness, ACID compliance, and excellent support for complex queries and concurrent transactions.",
    "### Flyway",
    "Flyway was integrated for database migration. It ensured that all database schema changes were version-controlled, reproducible, and seamlessly applied across different environments.",
    "### Maven",
    "Apache Maven was used for project management and build automation. It effectively handled dependency management and the project build lifecycle.",
    "### Postman",
    "Postman served as the primary tool for designing, documenting, and testing the REST APIs. It allowed the team to create comprehensive test collections to verify API endpoints independently of the frontend.",
    "### Redis",
    "Redis was utilized as an in-memory data structure store to facilitate caching, significantly improving the read performance of frequently accessed data."
]
add_chapter(document, "3", "Organization & Tools", ch3)

ch4 = [
    "The LMS backend was designed utilizing a strict RESTful architecture. The system exposes a series of stateless, resource-oriented APIs that the frontend application consumes. Communication between the client and server occurs via HTTP, with data exchanged predominantly in JSON format.",
    "The architectural design follows a classic layered pattern comprising the Controllers (handling HTTP requests and responses), Services (containing business logic), and Repositories (managing database interactions). To maintain a clean separation of concerns, Data Transfer Objects (DTOs) were utilized for client-server communication, preventing internal entity models from being directly exposed.",
    "[PLACEHOLDER] [Insert Figure 4.1: Workflow.png diagram from repository here]",
    "Figure 4.1 illustrates the technical workflow of the system. Incoming requests are authenticated and routed to the appropriate controller. The controller delegates processing to the service layer, which interacts with the PostgreSQL database via Spring Data JPA repositories. All database schema evolutions are managed sequentially by Flyway scripts executed during application startup."
]
add_chapter(document, "4", "Methodology / System Design", ch4)

ch5 = [
    "The implementation phase was structured around iterative, module-by-module API development. We began by establishing the core curriculum architecture. This involved designing schemas and APIs for hierarchical entities: Categories, Courses, Modules, Submodules, and Content. Each entity required robust CRUD (Create, Read, Update, Delete) operations with strict relational integrity.",
    "Following the curriculum module, we implemented the Assessments module, enabling the creation of quizzes and exams. Subsequently, the Events and Event Registrations modules were developed to manage live sessions and webinars. Finally, the Batches & Student Management module was created to group users and track their progress.",
    "Database migrations were strictly controlled using Flyway. We authored SQL scripts for every schema change (e.g., V1__Create_Course_Table.sql, V2__Add_Foreign_Keys.sql). This ensured that database structures were perfectly synchronized with the application codebase across all development instances.",
    "To handle data transformation between our JPA Entities and the DTOs exposed via the APIs, we utilized MapStruct. MapStruct is a code generator that greatly simplifies the implementation of mappings between Java bean types, reducing boilerplate code and improving performance compared to reflection-based mappers.",
    "A critical architectural achievement was the implementation of a multi-tenant data isolation approach. To support an enterprise context where multiple distinct organizations might use the LMS, we designed the database to securely partition data using a tenant identifier. This ensured that API requests were always scoped to the authenticated user's organization, preventing any cross-tenant data leakage."
]
add_chapter(document, "5", "Implementation", ch5)

ch6 = [
    "The development phase culminated in the successful deployment of the backend architecture. All APIs developed for Curriculum, Assessments, Events, and Student Management were fully integrated with the React-based frontend application.",
    "The integration process involved extensive collaboration with the frontend team to ensure that the API contracts were strictly adhered to and that the JSON payloads matched the UI components' expectations. The successful integration validated our headless architecture approach, demonstrating the seamless decoupling of the frontend presentation from the backend logic.",
    "The backend system is currently deployed live and supports the fully functional web application.",
    "[PLACEHOLDER] [Insert Screenshot 1: Live LMS Dashboard here]",
    "[PLACEHOLDER] [Insert Screenshot 2: Course Management Interface here]",
    "[PLACEHOLDER] [Insert Screenshot 3: Event Registration API integration here]",
    "The live application can be accessed at: https://xebia-portal-lms.vercel.app/. The performance of the backend APIs in the live environment has been highly stable, handling requests with minimal latency."
]
add_chapter(document, "6", "Results and Discussion", ch6)

ch7 = [
    "Ensuring the reliability and data integrity of the backend was paramount. My testing strategy encompassed multiple layers of verification.",
    "Extensive API testing was conducted using Postman. I developed comprehensive Postman collections that covered all endpoints, testing positive scenarios (valid data) and negative scenarios (invalid payloads, missing parameters, unauthorized access). Tests included automated assertions on HTTP status codes, response times, and JSON body structures.",
    "Manual database verification was performed consistently. After triggering API requests (via Postman or the frontend UI), I directly queried the PostgreSQL database to verify that the corresponding records were accurately created, updated, or deleted, and that foreign key constraints and cascading behaviors functioned correctly.",
    "To verify the scalability and performance of the APIs, load testing was conducted. I simulated concurrent traffic by injecting data for 50 users simultaneously using automated scripts. The objective was to monitor the application's behavior under load, checking for database connection pool exhaustion, transaction deadlocks, or significant degradation in response times. The system successfully processed the concurrent requests without errors, validating the robustness of the Spring Boot and PostgreSQL configuration."
]
add_chapter(document, "7", "Testing", ch7)

ch8 = [
    "The two-month summer internship at Xebia IT Architects Pvt Ltd. was a highly successful and enriching experience. I successfully contributed to the complete lifecycle of developing an enterprise-grade Learning Management System from scratch. By designing the PostgreSQL database schemas and building the RESTful APIs using Spring Boot, I gained practical, in-depth knowledge of modern backend architecture, database migrations, and team collaboration in an agile environment.",
    "While the current backend architecture is robust and functional, several enhancements are planned for the future scope of the project:",
    "1. Implementation of JWT and RBAC: The current mock header authentication system will be replaced with robust JSON Web Tokens (JWT) for secure stateless authentication. Additionally, a comprehensive Role-Based Access Control (RBAC) mechanism will be integrated to manage granular permissions for distinct user roles (e.g., Admin, Instructor, Student).",
    "2. Cloud Object Storage Integration: Local file uploads (such as course materials and user avatars) will be migrated to a Cloud Object Storage solution, such as AWS S3. This will improve scalability, reduce server storage constraints, and enhance file delivery speeds.",
    "3. CI/CD and Kubernetes Deployment: To automate the deployment lifecycle, Continuous Integration and Continuous Deployment (CI/CD) pipelines will be established. Ultimately, the application will be containerized and deployed using Kubernetes to ensure high availability, auto-scaling, and simplified management in a production cloud environment."
]
add_chapter(document, "8", "Conclusion & Future Scope", ch8)

refs = [
    "[1] M. Fowler, \"Microservices,\" martinfowler.com, 2014. [Online]. Available: https://martinfowler.com/articles/microservices.html.",
    "[2] C. Walls, Spring in Action, 6th ed. Shelter Island, NY: Manning Publications, 2022.",
    "[3] \"Spring Boot Reference Documentation,\" Spring, 2024. [Online]. Available: https://docs.spring.io/spring-boot/docs/current/reference/htmlsingle/.",
    "[4] \"Flyway by Redgate - Database Migrations Made Easy,\" Flyway, 2024. [Online]. Available: https://flywaydb.org/."
]
add_chapter(document, "References", "References", refs)

annexures = [
    "[PLACEHOLDER] [Insert Monthly Attendance Sheet Here]",
    "[PLACEHOLDER] [Insert Supervisor Feedback Form Here]"
]
add_chapter(document, "Annexures", "Annexures", annexures)

document.save(r'C:\Project\Xebia\Xebia_LMS_Internship_Report.docx')
print("Successfully generated report at C:\\Project\\Xebia\\Xebia_LMS_Internship_Report.docx")
